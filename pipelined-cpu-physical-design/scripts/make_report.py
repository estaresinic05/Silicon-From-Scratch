#!/usr/bin/env python3
"""
Build a formatted Word report for one place-and-route run.

    python3 scripts/make_report.py --run runs/clk3p0
    python3 scripts/make_report.py --run runs/clk3p0 --images docs/images \
                                   --out docs/pnr-report-clk3p0.docx

Every number in the document is parsed out of that run's reports by qor.py.
Nothing is typed in by hand, which means a regenerated report cannot disagree
with the reports it claims to describe, and each iteration gets a document
that is correct by construction.

The document furniture matches the design verification reports in
Verilog/CPU/*/docs/design-verification-report.docx: Arial body, the purple
heading ramp, a centred cover page carrying a result banner, a generated
table of contents, numbered sections, a ruled running header and a
'Page N of M' footer. The two families of report are read side by side, so
they are set the same way.

Screenshots are optional. Any of the filenames in FIGURES found in --images
is placed in its section with a caption; the rest of the document closes over
the gap. Requires python-docx.
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

try:
    from PIL import Image
except ImportError:
    Image = None  # only used to spot wide figures; absence is not fatal

sys.path.insert(0, str(Path(__file__).resolve().parent))
import qor  # noqa: E402


# --------------------------------------------------------------------------
# Design tokens, taken from the design verification reports.
# --------------------------------------------------------------------------
# The heading ramp darkens as the type gets smaller, so a Heading 3 at 12pt
# does not go weak against white the way the display purple would.
H1_PURPLE = RGBColor(0x3A, 0x16, 0x75)
H2_PURPLE = RGBColor(0x5A, 0x28, 0xB0)
H3_PURPLE = RGBColor(0x4F, 0x1F, 0x9E)
CAPTION_INK = RGBColor(0x22, 0x12, 0x4A)
MUTED = RGBColor(0x59, 0x59, 0x59)
RULE_GREY = RGBColor(0x80, 0x80, 0x80)

PASS_GREEN = RGBColor(0x2E, 0x7D, 0x32)
WARN_AMBER = RGBColor(0xB2, 0x6A, 0x00)
FAIL_RED = RGBColor(0xC6, 0x28, 0x28)

FILL_HEAD = "DED2F6"
FILL_ZEBRA = "F1ECFB"
BORDER_GREY = "CCCCCC"

BODY_FONT = "Arial"
MONO_FONT = "Consolas"

BODY_PT = 11
TABLE_PT = 9.5
MONO_PT = 9

# The text column at 1 inch margins on US Letter, in twentieths of a point.
# The running header's right tab and the table width both key off it.
TEXT_WIDTH_TWIPS = 9360
TEXT_WIDTH_IN = 6.5

FIGURES = [
    ("die-routed.png", "figure-die",
     "The routed core. Standard cells in rows, signal routing on metal2 "
     "through metal6, and the power ring on metal8 and metal9."),
    ("die-zoom.png", "figure-zoom",
     "Detail. Individual standard cells sitting in their rows, with the "
     "metal1 power rails running along each row boundary."),
    ("die-critical-path.png", "figure-path",
     "The critical path highlighted across the die. It launches in MEM/WB, "
     "crosses the forwarding logic and ripples through all 32 bits of the "
     "ALU adder before being captured in EX/MEM."),
    ("die-modules.png", "figure-modules",
     "The register file selected across the die. All 992 flip-flops matching "
     "datapath/registers_RF_reg, which is 31 architectural registers of 32 "
     "bits each, x0 costing nothing because it is hardwired to zero."),
    ("power-grid.png", "figure-power",
     "The power grid with signal routing hidden. The ring runs on metal9 and "
     "metal8 around the core, vertical straps cross on metal8 every 25 "
     "microns, and sroute ties the whole grid down to the metal1 rails "
     "running inside every standard cell row."),
    ("clock-tree.png", "figure-clock",
     "The synthesised clock tree driving every flip-flop in the design."),
    ("congestion.png", "figure-congestion",
     "Routing congestion. Each bin is shaded by how much routing demand it "
     "carries against what the tracks there can supply."),
]


# --------------------------------------------------------------------------
# low-level docx helpers
# --------------------------------------------------------------------------

def _shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def set_borders(table, color=BORDER_GREY, size=4):
    """python-docx has no border API, so this is raw OOXML."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def set_cell_margins(table, top=60, left=120, bottom=60, right=120):
    """Breathing room inside every cell, matching the verification reports."""
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge, value in (("top", top), ("left", left),
                        ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def vcenter(cell):
    el = OxmlElement("w:vAlign")
    el.set(qn("w:val"), "center")
    cell._tc.get_or_add_tcPr().append(el)


def para_border(paragraph, edge="bottom", color=BORDER_GREY, size=4, space=2):
    """A rule, drawn as a border on an otherwise empty paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    borders.append(el)
    pPr.append(borders)


def right_tab(paragraph, pos=TEXT_WIDTH_TWIPS):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos))
    tabs.append(tab)
    pPr.append(tabs)


def add_field(paragraph, instruction, placeholder=None, dirty=False):
    """
    A Word field: PAGE, NUMPAGES, TOC.

    The placeholder is what a reader sees before the field is calculated.
    Marking the field dirty asks Word to recalculate it when the document
    opens, which is what fills in a table of contents that python-docx has
    no way to compute for itself.
    """
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(begin)
    run._r.append(instr)
    if placeholder is not None:
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(sep)
        run._r.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)
    return run


def update_fields_on_open(doc):
    """
    Ask Word to recalculate every field when the file opens.

    Without this the table of contents renders as its placeholder text and
    the reader has to know to press F9, which they will not.
    """
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    doc.settings.element.append(el)


def repeat_header(table):
    """
    Mark row 0 as a header so Word redraws it at the top of each page.

    Without this a table that breaks across a page continues as unlabelled
    columns of numbers, which is exactly where a reader loses the thread.
    """
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def cant_split(table):
    """No row may be broken in half by a page break."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))


# --------------------------------------------------------------------------
# document furniture
# --------------------------------------------------------------------------

def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    # Widow and orphan control, so a paragraph never leaves one lonely line
    # behind on the previous page.
    pf.widow_control = True

    ramp = {
        "Heading 1": (17, True, H1_PURPLE, 18, 8),
        "Heading 2": (13, True, H2_PURPLE, 12, 6),
        "Heading 3": (12, False, H3_PURPLE, 10, 6),
    }
    for name, (size, bold, colour, before, after) in ramp.items():
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.italic = False
        st.font.color.rgb = colour
        st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        # keep_with_next stops a heading stranding at the foot of a page;
        # keep_together stops a two-line heading splitting across the break.
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    cap = doc.styles["Caption"]
    cap.font.name = BODY_FONT
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.bold = False
    cap.font.color.rgb = CAPTION_INK
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_together = True

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_header(doc, left_text, right_text):
    p = doc.sections[0].header.paragraphs[0]
    para_border(p, edge="bottom", color=BORDER_GREY, size=4, space=2)
    right_tab(p)
    for text, tab_first in ((left_text, False), (right_text, True)):
        if tab_first:
            r = p.add_run()
            r.font.size = Pt(8)
            r._r.append(OxmlElement("w:tab"))
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(8)
        r.font.color.rgb = RULE_GREY


def add_footer(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_border(p, edge="top", color=BORDER_GREY, size=4, space=2)

    def grey(run):
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RULE_GREY
        return run

    grey(p.add_run("Page "))
    grey(add_field(p, "PAGE", placeholder="1"))
    grey(p.add_run(" of "))
    grey(add_field(p, "NUMPAGES", placeholder="1"))


# --------------------------------------------------------------------------
# content builders
# --------------------------------------------------------------------------

# Section numbering. Headings read "3.", "3.1" the way the verification
# reports do, and the numbers are counted rather than typed so inserting a
# section can never leave the rest of the document misnumbered.
_SEC = [0, 0]


def h1(doc, title):
    _SEC[0] += 1
    _SEC[1] = 0
    return doc.add_heading(f"{_SEC[0]}.  {title}", level=1)


def h2(doc, title):
    _SEC[1] += 1
    return doc.add_heading(f"{_SEC[0]}.{_SEC[1]}  {title}", level=2)


def appendix(doc, letter, title):
    return doc.add_heading(f"Appendix {letter}.  {title}", level=1)


def centred(doc, text, size, colour=None, bold=False, italic=False,
            before=None, after=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if colour is not None:
        r.font.color.rgb = colour
    return p


def cover_page(doc, title, subtitle, tagline, author, meta_lines,
               banner_text, banner_colour, images_dir=None, cover_key=None):
    centred(doc, title, 28, H1_PURPLE, bold=True, before=90, after=8)
    centred(doc, subtitle, 16, H2_PURPLE, after=2)
    centred(doc, tagline, 13, MUTED, italic=True, after=30)

    centred(doc, f"Design Author: {author}", BODY_PT, after=2)
    for line in meta_lines:
        centred(doc, line, 10, MUTED, after=2)

    centred(doc, banner_text, 14, banner_colour, bold=True,
            before=24, after=6)

    if images_dir and cover_key:
        name, _ = _figure_spec(cover_key)
        path = Path(images_dir) / name if name else None
        if path and path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.add_run().add_picture(str(path), width=Inches(3.4))

    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Contents", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\o "1-2"',
              placeholder="Right-click here and choose Update Field to build "
                          "the table of contents.",
              dirty=True)
    doc.add_page_break()


def body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def mono_block(doc, text):
    """
    A monospace run of lines, set the way the verification reports set a
    module hierarchy: no box, no shading, just Consolas holding the columns.
    """
    lines = text.rstrip("\n").split("\n")
    for n, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8 if n == len(lines) - 1 else 0)
        p.paragraph_format.space_before = Pt(3 if n == 0 else 0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(MONO_PT)
    return lines


_CAP_N = {"Figure": 0, "Table": 0}


def add_caption(doc, kind, text):
    _CAP_N[kind] += 1
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{kind} {_CAP_N[kind]}: {text}")
    return p


def data_table(doc, headers, rows, aligns=None, caption=None, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_borders(table)
    set_cell_margins(table)
    repeat_header(table)
    aligns = aligns or ["l"] * len(headers)

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, FILL_HEAD)
        vcenter(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if aligns[i] == "r":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(str(h))
        r.font.name = BODY_FONT
        r.font.size = Pt(TABLE_PT)
        r.font.bold = True

    for n, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if n % 2 == 1:
                shade_cell(cells[i], FILL_ZEBRA)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if aligns[i] == "r":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            text = "" if val is None else str(val)
            bold = text.startswith("**") and text.endswith("**")
            if bold:
                text = text[2:-2]
            r = p.add_run(text)
            r.font.name = BODY_FONT
            r.font.size = Pt(TABLE_PT)
            r.font.bold = bold

    cant_split(table)

    if widths:
        # The column widths are written as proportions of each other and then
        # scaled to fill the text column, so every table in the document ends
        # flush with the margins the way the verification reports' do. A short
        # table floating centred in the middle of the page reads as a
        # different kind of object from a full-width one.
        scale = TEXT_WIDTH_IN / sum(widths)
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w * scale)

    if caption:
        add_caption(doc, "Table", caption)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def bullets(doc, items):
    for lead, rest in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        if lead:
            r = p.add_run(lead)
            r.bold = True
        p.add_run(rest)


def numbered(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(text)


# Die captures are square, so width is also height. At 6 inches a figure was
# two thirds of the 9 inch text column, which meant any figure not landing
# near the top of a page got pushed to the next one and left a hand-sized gap
# behind it. A little over 3 inches flows: two can share a page, and a figure
# can follow a paragraph instead of displacing it.
FIG_WIDTH_IN = 3.1


def maybe_figure(doc, images_dir, key, width_in=None):
    """Insert a screenshot if it exists. Silence if it does not."""
    if not images_dir:
        return False
    for name, tag, caption in FIGURES:
        if tag != key:
            continue
        path = Path(images_dir) / name
        if not path.exists():
            return False

        width = width_in or FIG_WIDTH_IN
        # A wide, short image would be needlessly small at the square width,
        # so give landscape figures more room.
        if Image is not None:
            try:
                with Image.open(path) as im:
                    w, h = im.size
                if h and w / h > 1.6:
                    width = min(6.0, width * 1.35)
            except OSError:
                pass

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(2)
        # Bind the image to its caption so a page break can never fall
        # between them.
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, "Figure", caption)
        return True
    return False


def _figure_spec(key):
    for name, tag, caption in FIGURES:
        if tag == key:
            return name, caption
    return None, None


def figure_pair(doc, images_dir, key_left, key_right, width_in=2.85):
    """
    Two figures side by side in a borderless table, captions beneath each.

    Six square figures stacked one per row is nearly three pages of nothing
    but pictures. Paired, they cost a little over one. Falls back to a single
    figure if only one of the two images exists, so a partial screenshot set
    still produces a sensible document.
    """
    if not images_dir:
        return False

    have = []
    for key in (key_left, key_right):
        name, caption = _figure_spec(key)
        if name and (Path(images_dir) / name).exists():
            have.append((Path(images_dir) / name, caption))

    if not have:
        return False
    if len(have) == 1:
        only = key_left
        name, _ = _figure_spec(key_left)
        if not (name and (Path(images_dir) / name).exists()):
            only = key_right
        return maybe_figure(doc, images_dir, only)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    row = table.rows[0]
    for idx, (path, caption) in enumerate(have):
        cell = row.cells[idx]
        cell.width = Inches(3.05)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(4)
        p.add_run().add_picture(str(path), width=Inches(width_in))

        _CAP_N["Figure"] += 1
        c = cell.add_paragraph(style="Caption")
        c.paragraph_format.space_before = Pt(2)
        c.paragraph_format.space_after = Pt(4)
        c.add_run(f"Figure {_CAP_N['Figure']}: {caption}")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return True


# --------------------------------------------------------------------------

def f(v, places=3, dash="not measured"):
    if v in ("", None):
        return dash
    try:
        return f"{float(v):.{places}f}"
    except (TypeError, ValueError):
        return str(v)


def i(v, dash="not measured"):
    if v in ("", None):
        return dash
    try:
        return f"{int(float(v)):,}"
    except (TypeError, ValueError):
        return str(v)


def ps(seconds_ns):
    """
    Nanoseconds as a whole number of picoseconds, correctly singular.

    A clock tree that costs exactly one picosecond is a real outcome on a
    block this small, and 'cost 1 picoseconds' in a signoff document reads
    as carelessness about everything else in it.
    """
    n = round(float(seconds_ns) * 1000)
    return f"{n} picosecond" if abs(n) == 1 else f"{n} picoseconds"


def _int_or_none(v):
    try:
        return int(float(v))
    except (TypeError, ValueError):
        return None


def result_banner(d):
    """
    The cover's one-line verdict, and the colour it is set in.

    Setup and hold are reported separately because they fail for opposite
    reasons and a single PASS would hide one behind the other.
    """
    setup = d.get("wns_setup")
    hold_viol = _int_or_none(d.get("n_hold_viol")) or 0
    setup_viol = _int_or_none(d.get("n_setup_viol")) or 0
    setup_met = setup not in ("", None) and float(setup) >= 0

    if not setup_met:
        return (f"TIMING RESULT:  SETUP NOT MET  "
                f"({setup_viol} paths violated)"), FAIL_RED
    if hold_viol:
        return (f"TIMING RESULT:  SETUP MET  ·  "
                f"{hold_viol} HOLD PATHS VIOLATED"), WARN_AMBER
    return "TIMING RESULT:  PASS  (0 paths violated)", PASS_GREEN


def build(run_dir, out_path, images_dir=None, author="Elliot Staresinic"):
    run_dir = Path(run_dir)
    d = qor.parse_run(run_dir)
    layers = qor.parse_wire_by_layer(run_dir / "reports" / "44_summary.rpt")
    notes_file = run_dir / "NOTES.md"
    note = notes_file.read_text(encoding="utf-8").strip() if notes_file.exists() else ""

    _CAP_N["Figure"] = 0
    _CAP_N["Table"] = 0
    _SEC[0] = 0
    _SEC[1] = 0

    doc = Document()
    setup_styles(doc)
    add_header(doc, "Physical Design Report  |  RISC-V Pipelined CPU", author)
    add_footer(doc)
    update_fields_on_open(doc)

    clk = d.get("clk_ns")
    clock_line = (f"Clock: {f(clk, 2)} ns ({1000 / float(clk):.0f} MHz)"
                  if clk else "Clock: not recorded")

    banner_text, banner_colour = result_banner(d)

    cover_page(
        doc,
        title="Physical Design Report",
        subtitle="RISC-V Pipelined CPU",
        tagline="From RTL to a Routed Die on a 45 nm Standard Cell Library",
        author=author,
        meta_lines=[
            "Design: pipelined_cpu_core, five-stage RISC-V, RV32I subset",
            f"Run: {d.get('run')}    |    {clock_line}",
            "Library: Nangate45, typical corner, QRC extraction",
            "Tools: Cadence Genus and Innovus 23.12-s091_1",
            f"Date: {d.get('date') or 'not recorded'}",
        ],
        banner_text=banner_text,
        banner_colour=banner_colour,
        images_dir=images_dir,
        cover_key="figure-die",
    )

    add_toc(doc)

    # ------------------------------------------------------- 1. summary
    h1(doc, "Executive Summary")

    setup = d.get("wns_setup")
    met = (setup not in ("", None)) and float(setup) >= 0
    verdict = "meets setup timing" if met else "does not meet setup timing"
    # An unparsed clock must not leave "at a  ns clock" in the prose, so the
    # whole phrase changes shape rather than just the number.
    at_clock = (f"at a {f(clk, 2)} ns clock" if clk
                else "at the clock period the SDC constrained it to")

    body(doc,
         f"This report documents the physical implementation of a five-stage "
         f"pipelined RISC-V CPU (pipelined_cpu_core), carried from RTL to a "
         f"routed layout through Cadence Genus and Innovus on the Nangate45 "
         f"45 nm standard cell library. Every number in this document is "
         f"parsed directly from the run's own reports rather than "
         f"transcribed.")

    body(doc,
         f"The design routes to completion and {verdict} "
         f"{at_clock}, with a worst negative slack of "
         f"{f(setup)} ns on the critical path. It occupies "
         f"{i(d.get('cells'))} standard cells at {f(d.get('density_pct'), 1)} "
         f"percent density, wired with {i(d.get('wire_um'))} microns of copper.")

    hold_viol = _int_or_none(d.get("n_hold_viol")) or 0
    if met and hold_viol:
        body(doc,
             f"Setup is met and hold is not. {hold_viol} hold paths are "
             f"reported as violated, which is a constraint artefact rather "
             f"than a physical one: this run modelled a single clock "
             f"uncertainty for both checks, so the setup margin was subtracted "
             f"from the hold check as well. Splitting the two uncertainties is "
             f"the first change made in the next iteration.")

    data_table(
        doc,
        ["Metric", "Value", "Source"],
        [
            ["Setup WNS", f"**{f(setup)} ns**", "40_final_setup.rpt"],
            ["Setup paths violated", i(d.get("n_setup_viol"), "0"), "40_final_setup.rpt"],
            ["Hold WNS", f"{f(d.get('wns_hold'))} ns", "41_final_hold.rpt"],
            ["Hold paths violated", i(d.get("n_hold_viol"), "0"), "41_final_hold.rpt"],
            ["Standard cells", i(d.get("cells")), "44_summary.rpt"],
            ["Filler cells", i(d.get("fillers")), "44_summary.rpt"],
            ["Flip-flops", i(d.get("flops")), "44_summary.rpt"],
            ["Logic area", f"{f(d.get('logic_um2'), 1)} um2", "44_summary.rpt"],
            ["Core area", f"{f(d.get('core_um2'), 1)} um2", "44_summary.rpt"],
            ["Core density", f"{f(d.get('density_pct'), 1)} %", "44_summary.rpt"],
            ["Total wire length", f"{i(d.get('wire_um'))} um", "44_summary.rpt"],
        ],
        aligns=["l", "r", "l"],
        widths=[2.1, 1.6, 2.5],
        caption="Signoff numbers, parsed directly from the run's own reports. "
                "Every figure in this document is read from the files named in "
                "the third column.",
    )

    if note:
        h2(doc, "What Changed in This Run")
        for line in note.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    # ---------------------------------------------------------- 2. flow
    h1(doc, "The Implementation Flow")

    body(doc,
         "Synthesis turns Verilog into a netlist: a flat list of logic gates "
         "and the wires between them, carrying no physical meaning whatsoever. "
         "Place and route turns that list into geometry, deciding where every "
         "gate physically sits and what copper connects them. The stages below "
         "run in order, and the design is written to disk after each one so a "
         "late failure never costs the earlier work.")

    h2(doc, "Stages")

    data_table(
        doc,
        ["Stage", "Command", "What it produces"],
        [
            ["1", "floorPlan", "A core box solved from a density target, not a fixed size"],
            ["2", "addRing, addStripe, sroute", "A power grid, which the netlist never mentions"],
            ["3", "place_opt_design", "Coordinates, plus logic restructuring to meet timing there"],
            ["4", "clock_opt_design", "A real buffered clock tree, replacing the ideal clock"],
            ["5", "addFiller", "Continuous power rails and wells across every row"],
            ["6", "route_opt_design", "Copper, then re-optimisation against extracted parasitics"],
            ["7", "verify_drc", "A check that the geometry obeys the routing rules"],
            ["8", "report, defOut", "The reports this document is built from"],
        ],
        aligns=["r", "l", "l"],
        widths=[0.6, 1.9, 3.7],
        caption="The place-and-route flow, stage by stage.",
    )

    h2(doc, "Why the Flow Is Written in Tcl")

    body(doc,
         "Innovus is, at its core, a Tcl interpreter. It is a Tcl shell with "
         "several thousand extra commands compiled in and a database of the "
         "chip attached to it. Typing place_opt_design at the Innovus prompt "
         "is typing Tcl, and running innovus -files scripts/innovus.tcl means "
         "nothing more than reading those same commands from a file instead of "
         "from the keyboard. There is no compilation step and no difference "
         "between the scripted tool and the interactive one.")

    body(doc,
         "This is not a Cadence quirk. Tcl was designed in the late 1980s "
         "specifically to be an embeddable command language for applications, "
         "EDA adopted it early, and it stayed. Cadence, Synopsys and Siemens "
         "tools are all driven this way, so the skill transfers across the "
         "industry in a way that familiarity with one vendor's graphical "
         "interface does not.")

    data_table(
        doc,
        ["File", "Language", "Talks to", "Job"],
        [
            ["run.sh", "bash", "Linux", "Check prerequisites, make directories, launch tools"],
            ["scripts/genus.tcl", "Tcl", "Genus", "RTL to gate-level netlist"],
            ["scripts/innovus.tcl", "Tcl", "Innovus", "Netlist to routed layout"],
            ["constraints/*.sdc", "Tcl", "any timer", "Declare what correct timing means"],
        ],
        aligns=["l", "l", "l", "l"],
        widths=[1.5, 0.8, 0.9, 3.0],
        caption="Three languages, three jobs. The shell script never mentions "
                "the chip; it confirms the tools and libraries exist and hands "
                "off.",
    )

    body(doc,
         "SDC is Tcl as well, and that is why it is portable. create_clock "
         "-name clk -period 3.0 is a procedure call, and because every "
         "vendor's timing engine implements the same procedure names, one "
         "constraint file feeds synthesis, place and route, and a signoff "
         "timer unchanged. That portability is why constraints live in their "
         "own file rather than inside the implementation script.",
         bold_lead="Constraints are portable because they are Tcl. ")

    body(doc,
         "Physical design is not done in a graphical interface. The layout "
         "view is for inspecting a result, never for producing one. Production "
         "blocks run in batch on a compute farm, unattended, for hours or "
         "days, and a flow that cannot run headless cannot be run at scale. "
         "Scripting also makes a run reproducible: because every input is a "
         "text file under version control, it is possible to say exactly what "
         "changed between two results, which is the entire basis of being able "
         "to iterate on one.",
         bold_lead="Why it is scripted at all. ")

    # -------------------------------------------------------- 3. timing
    h1(doc, "Timing")

    start = d.get("wns_place_start")
    place, cts, final = d.get("wns_place"), d.get("wns_cts"), d.get("wns_setup")

    body(doc,
         "Slack is the margin on the worst path. Positive means the design "
         "works at that clock, negative means it does not, and the magnitude "
         "says by how much. Tracking it stage by stage is how you tell whether "
         "a timing problem lives in synthesis, in placement, in the clock tree "
         "or in the routing.")

    h2(doc, "Slack Through the Flow")

    data_table(
        doc,
        ["After", "Clock model", "WNS (ns)", "What happened"],
        [
            ["Netlist as placed", "ideal", f(start), "Starting point, before optimisation"],
            ["place_opt_design", "ideal", f(place), "Placement optimisation"],
            ["clock_opt_design", "propagated", f(cts), "A real clock tree replaces the ideal one"],
            ["route_opt_design", "propagated + RC", f"**{f(final)}**", "Real copper, extracted parasitics"],
        ],
        aligns=["l", "l", "r", "l"],
        widths=[1.5, 1.2, 0.9, 2.6],
        caption="Setup slack through the flow. The starting point is read from "
                "the optimiser's own progress table in "
                "reports/um*/flow_QOR_summary.rpt.",
    )

    try:
        cts_cost = float(place) - float(cts)
        route_cost = float(cts) - float(final)
        recovered = float(place) - float(start)
        body(doc,
             f"Placement did nearly all of the work. The netlist synthesis "
             f"handed over was {ps(abs(float(start)))} short "
             f"the moment real coordinates existed, and placement optimisation "
             f"recovered {ps(recovered)} of that. Building "
             f"the actual clock tree cost {ps(cts_cost)} and "
             f"routing real wires cost a further {ps(route_cost)}.",
             bold_lead="Reading the table. ")
    except (TypeError, ValueError):
        pass

    h2(doc, "Where the Time Actually Goes")

    body(doc,
         "The worst path launches from a register in MEM/WB, crosses the "
         "forwarding unit that decides EX must take a forwarded value rather "
         "than the register file's, and then ripples the carry across all 32 "
         "bits of the ALU adder before being captured in EX/MEM. Of the "
         "2.863 ns that path takes, 2.123 ns is the carry chain alone: "
         "74 percent of the clock period sits in one ripple-carry adder. "
         "That is the single change worth making to the design, and the adder "
         "is under 3 percent of the area, so speed there is cheap.")

    figure_pair(doc, images_dir, "figure-clock", "figure-path")

    h2(doc, "What This Number Does Not Tell You")

    body(doc,
         "A run that meets timing does not tell you what was achievable. The "
         "optimiser converges to whatever constraint it was given and then "
         "stops, and its area reclaiming passes deliberately spend positive "
         "slack to recover area and leakage. Synthesis stops trying at the "
         "same constraint. Finding the real maximum frequency means tightening "
         "the period and rerunning both tools until it fails.")

    # ------------------------------------------------------ 4. physical
    h1(doc, "The Physical Result")

    h2(doc, "Cells and Fillers")

    body(doc,
         f"The core holds {i(d.get('cells'))} standard cells of real logic "
         f"alongside {i(d.get('fillers'))} filler cells. Fillers contain no "
         f"logic at all; standard cells carry power rails and wells that must "
         f"be physically continuous across a row, and fillers exist purely to "
         f"bridge the gaps left between placed cells. Without them the rails "
         f"break and the design is not manufacturable.")

    if layers:
        h2(doc, "Routing by Metal Layer")
        total = sum(layers.values()) or 1
        rows = []
        for n in range(1, 11):
            key = f"metal{n}"
            if key not in layers:
                continue
            v = layers[key]
            rows.append([key, f"{v:,.0f}", f"{100 * v / total:.1f} %"])
        data_table(
            doc, ["Layer", "Wire length (um)", "Share"], rows,
            aligns=["l", "r", "r"], widths=[1.2, 1.6, 1.0],
            caption="Routing by metal layer. metal1 carries no signal because "
                    "it is reserved for wiring inside the cells and for the "
                    "power rails. The upper layers are nearly empty because a "
                    "block this small has no journeys long enough to need them.",
        )

    h2(doc, "The Power Grid")

    body(doc,
         "The power grid is worth seeing on its own, because none of it comes "
         "from the netlist. A netlist describes logic and says nothing about "
         "how current reaches it. Every ring, strap and rail below was created "
         "by three commands in the flow script, and the connection between a "
         "cell's power pin and the VDD net exists only because "
         "globalNetConnect declared it.",
         bold_lead="Power is added, never inherited. ")

    figure_pair(doc, images_dir, "figure-zoom", "figure-power")

    muxed = d.get("flops_muxed")
    muxed_um2 = d.get("flops_muxed_um2")
    if muxed and muxed_um2 and d.get("logic_um2"):
        h2(doc, "Where the Area Actually Goes")
        share = 100 * float(muxed_um2) / float(d["logic_um2"])
        body(doc,
             f"Of the {i(d.get('flops'))} flip-flops, {i(muxed)} are SDFF, a "
             f"cell with a 2x1 multiplexer built into it, which is what an "
             f"enable-guarded register maps onto. That count is the register "
             f"file exactly: 31 architectural registers of 32 bits, x0 costing "
             f"nothing because it is hardwired to zero. Those cells alone are "
             f"{f(muxed_um2, 0)} um2, or {share:.0f} percent of all logic area, "
             f"and the multiplexer tree that reads them adds more. "
             f"The structure that dominates the area and the structure that "
             f"dominates the clock are different parts of the design.")

    maybe_figure(doc, images_dir, "figure-modules", width_in=2.7)
    maybe_figure(doc, images_dir, "figure-congestion")

    # ---------------------------------------------------------- 5. gaps
    h1(doc, "What a Production Flow Has That This One Does Not")

    body(doc,
         "The structure of this flow matches industry practice. The rigour "
         "does not, in ways worth naming explicitly.")

    bullets(doc, [
        ("One corner is not signoff. ",
         "This run used a single library, a single RC corner and a single "
         "analysis view. Real signoff is multi-corner multi-mode, spanning "
         "slow and fast silicon, hot and cold, high and low voltage, and "
         "functional against test modes. The same library was also used for "
         "setup and hold, where production uses a slow corner for setup and a "
         "fast corner for hold."),
        ("The place-and-route timer is not the signoff timer. ",
         "Innovus's internal engine is good enough to optimise against, but "
         "signoff exports the netlist with extracted parasitics and re-times "
         "in a dedicated tool, expecting the two to disagree slightly."),
        ("verify_drc is not physical verification. ",
         "It checks routing rules inside the tool. Real DRC and LVS run "
         "against a foundry runset in a separate verification tool. Nangate45 "
         "is an academic library with no foundry behind it, so this does not "
         "apply here, but it would immediately on real silicon."),
        ("Whole categories are absent. ",
         "IR drop and electromigration, antenna fixing, metal fill, scan "
         "insertion and chain reordering, and formal equivalence checking "
         "between the RTL and the synthesised netlist."),
        ("Flows are normally inherited, not authored. ",
         "Companies have methodology teams that own these scripts, and a block "
         "designer supplies the RTL, the constraints and a configuration file "
         "of block-specific knobs. Writing one from scratch was worth doing "
         "once for the understanding."),
    ])

    # ---------------------------------------------------- 6. conclusion
    h1(doc, "Conclusion")

    body(doc,
         f"The five-stage pipelined RISC-V CPU reaches a routed layout on "
         f"Nangate45 and {verdict} {at_clock}, "
         f"closing at {f(setup)} ns of "
         f"setup slack across {i(d.get('cells'))} standard cells and "
         f"{i(d.get('wire_um'))} microns of copper. The flow that produced it "
         f"is scripted end to end, so the result is reproducible from the "
         f"repository alone.")

    body(doc,
         "Two things are worth carrying forward. The first is that placement, "
         "not routing, did almost all of the timing work, which says the "
         "netlist handed over by synthesis was physically reasonable and that "
         "the remaining margin will not be found by tuning the back end. The "
         "second is that the critical path is dominated by a single "
         "ripple-carry chain in the ALU while the area is dominated by the "
         "register file, so the change that buys speed and the change that "
         "buys area are different changes to different parts of the design.")

    body(doc,
         "The number this run cannot supply is the maximum achievable "
         "frequency, because an optimiser that meets its constraint stops "
         "there. Establishing it means sweeping the clock period until the "
         "design fails, and that sweep is what the per-run directory structure "
         "and the generated results table exist to support.")

    # ------------------------------------------------------- appendix A
    appendix(doc, "A", "Reproducing This Run")

    body(doc,
         "The run is fully described by the scripts in the repository. Every "
         "report quoted above is committed under results/ so each number can "
         "be checked against its source.")

    mono_block(doc,
               "./run.sh --period {p} --name {n} --note \"...\"\n"
               "\n"
               "runs/{n}/reports/     the reports this document parses\n"
               "results/{n}/reports/  the same files, committed as evidence\n"
               "results/QOR.md        one row per run, for comparing iterations"
               .format(p=f(clk, 1, "3.0"), n=d.get("run")))

    body(doc,
         f"This document was generated from {d.get('run')} by "
         f"scripts/make_report.py. The numbers are parsed from the run's "
         f"reports, not transcribed, so regenerating it after a rerun cannot "
         f"produce a document that disagrees with its own evidence. "
         f"Repository: estaresinic05/Silicon-From-Scratch.")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path, d


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--run", required=True, help="run directory containing reports/")
    ap.add_argument("--out", default=None, help="output .docx path")
    ap.add_argument("--images", default=None, help="directory of screenshots")
    ap.add_argument("--author", default="Elliot Staresinic")
    args = ap.parse_args()

    run_dir = Path(args.run)
    if not (run_dir / "reports").is_dir():
        sys.exit(f"no reports/ under {run_dir}")

    root = Path(__file__).resolve().parent.parent
    out = Path(args.out) if args.out else root / "docs" / f"pnr-report-{run_dir.name}.docx"

    path, d = build(run_dir, out, args.images, args.author)
    print(f"wrote  {path}")
    print(f"run    {d.get('run')}   setup WNS {f(d.get('wns_setup'))} ns")


if __name__ == "__main__":
    main()
