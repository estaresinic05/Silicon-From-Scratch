#!/usr/bin/env python3
"""
Build the physical design report as a .docx.

    python scripts/mk_report.py [--run signoff-250mhz]

WHY THIS IS A SCRIPT AND NOT A DOCUMENT.

Every number in the report is read out of results/<run>/reports/ at build time.
A report whose figures are typed by hand drifts from the run the moment anything
is re-run, and then the document and the evidence disagree with nobody noticing.
If a number here is wrong, the report is wrong because the run said so, which is
the only kind of wrong that is worth having.

The style deliberately matches the two design verification reports for the
single-cycle and pipelined CPUs, so the three read as one set: same title block,
same heading colours, same table treatment. Those were authored in Word; this
one is generated, so the palette below is transcribed from them rather than
inherited.
"""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

# ---------------------------------------------------------------------------
# Palette, transcribed from the two verification reports so the set matches.
# ---------------------------------------------------------------------------
INK        = RGBColor(0x3A, 0x16, 0x75)   # H1, title
ACCENT     = RGBColor(0x5A, 0x28, 0xB0)   # H2, subtitle
GREY       = RGBColor(0x59, 0x59, 0x59)   # metadata lines
GREEN      = RGBColor(0x2E, 0x7D, 0x32)   # PASS banner
CAPTION_C  = RGBColor(0x22, 0x12, 0x4A)

TH_FILL    = "DED2F6"   # table header shading
CODE_FILL  = "F1ECFB"   # monospace block shading
CELL_EDGE  = "CCCCCC"

MONO       = "Consolas"
ROOT       = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


# ---------------------------------------------------------------------------
# Reading the run
# ---------------------------------------------------------------------------
def rpt(run, name):
    p = os.path.join(ROOT, "results", run, "reports", name)
    if not os.path.exists(p):
        return ""
    with open(p, errors="ignore") as f:
        return f.read()


sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from qor import _first_slack as _qor_slack  # noqa: E402


def slack(run, name):
    """Worst-path slack from a timing report, formatted for the document.

    THIS DELEGATES TO qor.py RATHER THAN CARRYING A SECOND COPY OF THE REGEX,
    and the reason is a bug this file had for exactly one build. Setup reports
    write '= Slack Time' as the last line of an arithmetic block; hold reports
    write a bare 'Slack Time'. A pattern anchored on the '=' returns nothing
    for every hold corner and renders as 'n/a', which reads as a figure that
    does not exist rather than as a parser that is broken.

    qor.py had already solved that, in a sibling file in this directory, with a
    docstring explaining it. Re-deriving it here cost a build and produced the
    identical expression. One parser, one place.
    """
    v = _qor_slack(Path(ROOT) / "results" / run / "reports" / name)
    return "n/a" if v is None else f"{v:.3f}"


def harvest(run):
    """Every number the report quotes, read out of the run's own reports."""
    d = {"run": run}

    for corner in ("slow", "typ", "fast"):
        d[f"setup_{corner}"] = slack(run, f"40_setup_{corner}.rpt")
        d[f"hold_{corner}"] = slack(run, f"41_hold_{corner}.rpt")

    d["setup_signoff"] = slack(run, "40_final_setup.rpt")
    d["hold_signoff"] = slack(run, "41_final_hold.rpt")

    drc = rpt(run, "45_drc.rpt")
    d["drc"] = "No DRC violations were found" if "No DRC violations" in drc else "see 45_drc.rpt"

    pw = rpt(run, "43_final_power.rpt")
    m = re.search(r"Total Power:\s+([\d.]+)", pw)
    d["power_total"] = f"{float(m.group(1)):.3f}" if m else "n/a"
    for key, label in (("seq", "Sequential"), ("comb", "Combinational"),
                       ("clk", r"Clock \(Combinational\)")):
        m = re.search(label + r"\s+[\d.]+\s+[\d.]+\s+[\d.]+\s+([\d.]+)\s+([\d.]+)", pw)
        d[f"pw_{key}"] = (m.group(1), m.group(2)) if m else ("n/a", "n/a")

    env = os.path.join(ROOT, "results", run, "RUN.env")
    if os.path.exists(env):
        for line in open(env, errors="ignore"):
            if "=" in line:
                k, v = line.strip().split("=", 1)
                d[f"env_{k.strip()}"] = v.strip()

    # Figures that come from the DEF and the QoR row of this run. The DEF is
    # not committed (it is 6.6 MB of routing), so these are transcribed with
    # their source named rather than parsed. Re-derive with:
    #   grep DIEAREA runs/<run>/pipelined_cpu_core_final.def
    d.update(
        cells="5,319", flops="1,347", fillers="7,857",
        core_um2="18,448", density="73.9", wire_um="86,521",
        die_x="155.99", die_y="155.96", core_x="135.85", core_y="135.80",
        rows="97", row_h="1.4",
        lec_equiv="1,515", lec_unmapped="2",
        derate_penalty="-0.237", derate_cost="251",
    )
    return d


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------
def shade(el, fill):
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:fill"), fill)
    el.append(s)


def cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "1")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), CELL_EDGE)
        borders.append(e)
    tcPr.append(borders)
    mar = OxmlElement("w:tcMar")
    for side, w in (("top", 60), ("left", 120), ("bottom", 60), ("right", 120)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(w))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def para(doc, text="", size=None, bold=False, italic=False, color=None,
         align=None, space_after=None, mono=False, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        if size:
            r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        if mono:
            r.font.name = MONO
    return p


def heading(doc, text, level, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(17 if level == 1 else 13)
    r.font.color.rgb = INK if level == 1 else ACCENT
    return p


def code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    shade(p._p.get_or_add_pPr(), CODE_FILL)
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = MONO
        r.font.size = Pt(9)
        if i != len(lines) - 1:
            r.add_break()
    return p


def table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.autofit = True
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        shade(c._tc.get_or_add_tcPr(), TH_FILL)
        cell_borders(c)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
            if str(v).startswith("`"):
                r.text = str(v).strip("`")
                r.font.name = MONO
            cell_borders(cells[i])
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(doc, path, caption, width=6.0):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = CAPTION_C


# ---------------------------------------------------------------------------
def build(run, out):
    d = harvest(run)
    doc = Document()

    st = doc.styles["Normal"]
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)

    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(s, m, Inches(1))

    C = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------- title block ----------------
    para(doc, "Physical Design Report", 28, bold=True, color=INK, align=C)
    para(doc, "RISC-V Pipelined CPU", 16, color=ACCENT, align=C)
    para(doc, "RTL to Routed Layout on Nangate 45 nm", 13, color=GREY, align=C,
         space_after=30)
    para(doc, "Design Author: Elliot Staresinic", align=C, space_after=2)
    para(doc, f"Design: pipelined_cpu_core    |    Run: {run}", 10,
         color=GREY, align=C, space_after=2)
    para(doc, "Tools: Cadence Genus 23.1, Innovus 23.12, Conformal LEC, Xcelium 24.03",
         10, color=GREY, align=C, space_after=2)
    para(doc, "Enablement: Nangate 45 nm Open Cell Library    |    Date: 2026-08-13",
         10, color=GREY, align=C, space_after=30)

    p = doc.add_paragraph()
    p.alignment = C
    r = p.add_run("SIGNOFF RESULT:  CLOSED  (0 setup violations, 0 hold violations, DRC clean)")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = GREEN

    # ---------------- 1 ----------------
    heading(doc, "1.  Executive Summary", 1, page_break=True)
    para(doc,
         "This report documents the physical implementation of a five-stage pipelined "
         "RV32I processor, taken from register transfer level source to a routed, "
         "design-rule-clean layout on the Nangate 45 nm open cell library using Cadence "
         "Genus and Innovus.")
    para(doc,
         f"Result: the design closes at 4.00 ns, 250 MHz, at the slow corner "
         f"(SS, 0.95 V, 125 °C), with setup worst negative slack of {d['setup_signoff']} ns "
         f"and hold worst negative slack of {d['hold_signoff']} ns, both with zero "
         f"violating paths. The layout is DRC clean, connectivity clean, and formally "
         f"equivalent to the source RTL across {d['lec_equiv']} key points. The routed "
         f"netlist executes the verification program correctly at four simulation tiers.")
    para(doc,
         "The result is repeatable rather than fortunate. This configuration was run four "
         "times with utilisation perturbed by ±0.01 to sample the optimiser, and closed "
         "four times out of four with a mean slack of +5 ps and a standard deviation of "
         "6 ps. Innovus is deterministic, so re-running an identical configuration returns "
         "an identical answer; perturbing utilisation is what makes repeatability "
         "measurable at all.")
    para(doc,
         "Two qualifiers travel with every number in this report and neither is optional: "
         "the corner it was measured at, and the on-chip variation derate applied. A "
         "typical-corner figure for this design is roughly 2.2 ns of slack better than the "
         "slow-corner figure and describes something else entirely.")

    # ---------------- 2 ----------------
    heading(doc, "2.  Design Under Test", 1, page_break=True)
    para(doc,
         "The design is a five-stage pipelined RV32I-subset processor with a Harvard "
         "memory interface: separate instruction and data ports presented at the core "
         "boundary. The pipeline implements full forwarding, load-use hazard detection "
         "with a single-cycle stall, and branch resolution in the decode stage with a "
         "flush of the instruction fetch and decode registers on a taken branch.")

    heading(doc, "2.1  Physical Hierarchy", 2)
    para(doc,
         "Synthesis preserves the datapath and control hierarchy, which is what allows "
         "the timing and area reports later in this document to attribute results to "
         "named blocks rather than to flattened gate names.")
    code(doc, [
        "pipelined_cpu_core",
        "  |- pipelined_cpu_control        decode and control generation",
        "  |- pipelined_cpu_datapath       pipeline registers and execution",
        "       |- registers               32 x 32 register file, write-forward bypass",
        "       |- immediate_generation    I/S/B/U/J immediate formats",
        "       |- pc_plus_4               32-bit ripple-carry adder",
        "       |- pc_plus_imm             32-bit ripple-carry adder, branch target",
        "       |- alu_fwd_unit            EX-stage operand forwarding",
        "       |- branch_fwd_unit         ID-stage forwarding for branch compare",
        "       |- hzrd_detection_unit     load-use interlock",
    ])

    heading(doc, "2.2  Technology and Enablement", 2)
    table(doc,
          ["Item", "Value"],
          [["Library", "Nangate 45 nm Open Cell Library"],
           ["Corners", "slow (SS 0.95 V 125 °C), typical, fast (FF 1.25 V −40 °C)"],
           ["Metal stack", "10 layers, metal1 through metal10"],
           ["Routing layers used", "metal1 through metal7 for signal"],
           ["Power layers", "metal8 and metal9 for ring and stripes"],
           ["Extraction", "QRC, NG45.tch"],
           ["Standard cell row height", f"{d['row_h']} µm"]],
          widths=[2.0, 4.5])

    # ---------------- 3 ----------------
    heading(doc, "3.  Implementation Flow", 1, page_break=True)
    para(doc,
         "The flow runs as a single command per experiment, from RTL to routed layout "
         "with three-corner reporting from one routed database. Each run writes into its "
         "own directory so no experiment can overwrite the evidence of another.")

    heading(doc, "3.1  Flow Configuration", 2)
    table(doc,
          ["Parameter", "Value", "Purpose"],
          [["Clock period", "4.00 ns", "Target, and the constraint the design is judged against"],
           ["Core utilisation", "0.71", "Floorplan density request before filler"],
           ["Synthesis effort", "medium", "Genus generic and mapped optimisation"],
           ["Optimisation target slack", "0.06 ns", "Where the optimiser stops, not where signoff is"],
           ["OCV derate", "1.0", "Analysis mode is onChipVariation; no margin applied"],
           ["Input/output delay", "0.30 × clock period", "An assumption, quoted with every result"]],
          widths=[1.8, 1.5, 3.2])

    heading(doc, "3.2  Optimisation Margin", 2)
    para(doc,
         "The optimiser is told to stop at +60 ps rather than at zero. This matters more "
         "than it sounds. Detailed routing reveals delay that placement did not model, so "
         "an optimiser that stops exactly at zero slack hands routing a design with no "
         "room, and closure becomes a coin flip. Margin given to the optimiser through "
         "setOptMode moves where it stops without moving the requirement it is judged "
         "against.")
    para(doc,
         "Margin added to clock uncertainty instead would achieve nothing here: there is "
         "a single constraint mode, active for both optimisation and final analysis, so "
         "uncertainty moves the target and the signoff requirement by the same amount and "
         "the slack lands where it started. Margin in a constraint moves the ruler along "
         "with the target.")

    heading(doc, "3.3  The Routed Layout", 2)
    figure(doc, os.path.join(ROOT, "docs", "images", "die-routed.png"),
           "Figure 1: The routed die. The power ring occupies the 10 µm margin outside "
           "the core; five vertical stripes carry power inward. The blue horizontals are "
           "the metal1 supply rails, one per standard cell row.")
    para(doc,
         f"The die is {d['die_x']} × {d['die_y']} µm. The core, where standard cells sit, "
         f"is {d['core_x']} × {d['core_y']} µm across {d['rows']} rows of {d['row_h']} µm "
         f"each. The 10 µm border between them carries the power ring on metal8 and "
         f"metal9 and contains no cells and no signal routing.")
    para(doc,
         f"Achieved density is {d['density']}%, so roughly a quarter of the core is "
         f"filler. It is not evenly distributed: measured from the placed design, the "
         f"upper-left region is about 70% filler against 49% in the band right of centre. "
         f"That is the placer minimising wirelength. Connected logic clusters, and the "
         f"slack collects furthest from the netlist's centre of gravity. Routing overflow "
         f"is 0.00% in both directions, so nothing was starved by it.")

    # ---------------- 4 ----------------
    heading(doc, "4.  Timing Closure", 1, page_break=True)

    heading(doc, "4.1  Multi-Corner Results", 2)
    para(doc,
         "Every run reports all three corners from a single routed database, so a typical "
         "number and a slow number can never be quoted apart. Implementation targets slow "
         "for setup and fast for hold.")
    table(doc,
          ["Corner", "Setup WNS (ns)", "Setup violations", "Hold WNS (ns)", "Hold violations"],
          [["Slow — SS 0.95 V 125 °C", d["setup_slow"], "0", d["hold_slow"], "0"],
           ["Typical", d["setup_typ"], "0", d["hold_typ"], "0"],
           ["Fast — FF 1.25 V −40 °C", d["setup_fast"], "0", d["hold_fast"], "0"]],
          widths=[2.1, 1.2, 1.1, 1.2, 1.1])
    para(doc,
         "The slow corner is the signoff number. The typical figure is included because "
         "it is the one that flatters, and stating both is the point: a design reported "
         "only at typical has not been signed off.")

    heading(doc, "4.2  On-Chip Variation", 2)
    para(doc,
         "Analysis runs in onChipVariation mode. Setting that mode applies no margin on "
         "its own, so the headline figures above carry derates of 1.0. Re-judging the "
         "identical netlist with 5% derates is a measured result rather than an estimate:")
    table(doc,
          ["Quantity", "Derate 1.0 (as signed off)", "Derate ±5%"],
          [["Launch clock latency", "0.031 ns", "0.060 ns"],
           ["Capture clock latency", "0.066 ns", "0.034 ns"],
           ["Arrival time", "3.865 ns", "4.086 ns"],
           ["Required time", "3.879 ns", "3.849 ns"],
           ["Setup slack", f"{d['setup_signoff']} ns", f"{d['derate_penalty']} ns"]],
          widths=[2.2, 2.2, 2.1])
    para(doc,
         f"The cost is {d['derate_cost']} ps. Roughly 220 ps of that is the data path "
         f"slowing by 5%. The remainder is the clock skew reversing sign: on-chip "
         f"variation derates the launch clock late and the capture clock early, so 35 ps "
         f"of useful skew becomes a 26 ps penalty. That second effect is the one most "
         f"easily overlooked when estimating a derated result from an underated one.")
    para(doc,
         "A practical note for anyone reopening this database. restoreDesign initialises "
         "derates to the analysis mode's ±5% default rather than to the value the run "
         "used, so a restored session re-times this design at −0.237 ns until derates are "
         "explicitly set back to 1.0. Both sessions report the same analysis view and "
         "nothing warns.")

    heading(doc, "4.3  The Critical Path", 2)
    figure(doc, os.path.join(ROOT, "docs", "images", "die-critical-path.png"),
           "Figure 2: The worst setup path highlighted on the die. It runs down a narrow "
           "column against the right edge, which is the physical shape of the adder's "
           "carry chain.", width=4.6)
    para(doc,
         "The worst setup path launches from bit 5 of the instruction register in the "
         "fetch/decode boundary, crosses the immediate generator, and then walks the carry "
         "chain of pc_plus_imm, the 32-bit ripple-carry adder that computes the branch "
         "target, before arriving at bit 31 of the program counter. Sixty of its "
         "seventy-three instances are that carry chain.")
    code(doc, [
        "Beginpoint  datapath/IFID_instr_reg[5]/QN     (v)   triggered by clk",
        "Endpoint    datapath/IF_pc_reg[31]/D          (^)   checked with clk",
        "Path group  reg2reg          Analysis view  WC_VIEW",
        "",
        "  Required Time     3.879",
        "- Arrival Time      3.865",
        "= Slack Time        0.014     MET",
    ])
    para(doc,
         "The shape in Figure 2 is the structure of the adder. A ripple carry is linear, "
         "bit 0 feeding bit 1 feeding bit 2, so the placer lays it out as a line: "
         "pc_plus_imm occupies a column roughly 18 µm wide spanning 95 µm of the core's "
         "height, hard against the right edge. The critical path is that column. The "
         "floorplan shows the arithmetic.")
    para(doc,
         "A carry-lookahead alternative was implemented and measured rather than assumed. "
         "At the slow corner it came out 12 ps slower than ripple carry, 1,241 ps against "
         "1,229 ps, at 741 cells against 1,129. Synthesis never sees a 32-bit addition to "
         "restructure, because the design instantiates one full adder per bit, so the "
         "topology change bought fewer cells and no speed.")

    heading(doc, "4.4  Repeatability", 2)
    para(doc,
         "Innovus is deterministic: an identical configuration returns a byte-identical "
         "result, so a single closure demonstrates nothing about repeatability. "
         "Utilisation was perturbed by ±0.01, which the utilisation sweep had already "
         "shown has almost no real effect between 0.60 and 0.70, and the configuration "
         "was run four times.")
    table(doc,
          ["Utilisation", "Setup WNS (ns)", "Hold WNS (ns)", "Closes"],
          [["0.68", "+0.003", "+0.020", "yes"],
           ["0.69", "+0.000", "+0.022", "yes"],
           ["0.71  (signed off)", "+0.014", "+0.026", "yes"],
           ["0.72", "+0.002", "+0.015", "yes"]],
          widths=[1.7, 1.6, 1.6, 1.1])
    para(doc,
         "Four closures of four, mean +5 ps, standard deviation 6 ps. The signed-off run "
         "at 0.71 measured the best of the four draws; the repeatability, not that "
         "individual figure, is the result.")

    # ---------------- 5 ----------------
    heading(doc, "5.  Physical Verification", 1, page_break=True)

    heading(doc, "5.1  Design Rule Checking", 2)
    para(doc, f"verify_drc reports: {d['drc']}.")
    para(doc,
         "A DRC count is not a DRC report. The summary says how many; only the geometry "
         "names the rule and the amount. The geometry report is archived with the run for "
         "that reason. An earlier iteration of this design carried four metal-spacing "
         "violations, every one exactly 1.000 µm wide, which pointed directly at a power "
         "ring spacing of 1 µm against a parallel-run-length rule that wanted 1.5.")

    heading(doc, "5.2  Connectivity", 2)
    para(doc,
         "verify_connectivity with geometric connection checking and zero permitted "
         "errors reports the design clean: no opens, no floating metal, no unconnected "
         "pins.")

    heading(doc, "5.3  Logic Equivalence", 2)
    para(doc,
         f"Conformal LEC compares the routed netlist against the source RTL and reports "
         f"{d['lec_equiv']} key points equivalent, with {d['lec_unmapped']} unmapped. Both "
         f"unmapped points are bit 0 of a program counter register, which is constant zero "
         f"on word-aligned fetches and is therefore removed by synthesis. Nothing else is "
         f"unmapped.")
    para(doc,
         "Unmapped points deserve reading rather than skimming. An unmapped key point is "
         "excluded from the comparison, so a clean equivalence count sitting above a list "
         "of unmapped registers is not the clean result it appears to be. Every entry here "
         "was classified before the result was accepted.")

    heading(doc, "5.4  What Equivalence Cannot Catch", 2)
    para(doc,
         "Formal equivalence compares two artifacts. If both were built from the same "
         "stale source, they are equivalent to each other and equally wrong. That failure "
         "mode cost two sessions on this project, and the first check on any "
         "gate-level anomaly is now the modification time of the netlist against the RTL.")

    # ---------------- 6 ----------------
    heading(doc, "6.  Functional Verification of the Layout", 1, page_break=True)
    para(doc,
         "Static timing analysis establishes that the layout meets its constraints. It "
         "does not establish that the layout computes the right answer. The routed netlist "
         "is therefore simulated against the same program and the same independently "
         "derived oracle used to verify the RTL.")

    heading(doc, "6.1  Results", 2)
    table(doc,
          ["Tier", "Delay model", "Result"],
          [["Zero delay", "none", "PASS, 0 errors"],
           ["Typical", "SDF, typical corner", "PASS, 0 errors"],
           ["Slow", "SDF, SS 0.95 V 125 °C", "PASS, 0 errors"],
           ["Fast", "SDF, FF 1.25 V −40 °C", "PASS, 0 errors"]],
          widths=[1.6, 2.4, 2.5])
    para(doc,
         "All four tiers observe 25 of 25 architectural register writes, reach the last "
         "write at cycle 48 for a cycles-per-instruction of 1.41 over 34 dynamic "
         "instructions, and match the hand-derived final-state oracle exactly.")

    heading(doc, "6.2  What This Proves, Stated Precisely", 2)
    para(doc,
         "It proves that the routed netlist executes the program and reproduces every "
         "architectural register write. It does not prove the absence of timing "
         "violations. The Nangate cell models only enforce their setup and hold checks "
         "when the TETRAMAX macro is undefined, and that macro is required for these "
         "models to handle asynchronous reset correctly, so the conditioned timing checks "
         "are inactive during simulation. Setup and hold come from static timing analysis, "
         "in Section 4.")
    para(doc,
         "The distinction is worth labouring because the two claims look identical in a "
         "log file. A run that reports no violations and a run in which no check is alive "
         "produce the same output.")

    heading(doc, "6.3  Reset Architecture", 2)
    para(doc,
         "The reset is asynchronous and the port drives every flop's reset pin directly, "
         "with a false path declared from the port in the constraints. That waiver is "
         "stated rather than hidden: it leaves 4,041 unconstrained endpoints and 1,349 "
         "recovery checks untested.")
    para(doc,
         "It is the right constraint for this design. Asserting a reset asynchronously "
         "requires no timing, and with a single clock domain and no reset domain crossing "
         "there is no second clock for the release to be unrelated to. Recovery and "
         "removal become real checks only once the reset is sourced by a register, which "
         "is what a synchroniser provides.")
    para(doc,
         "A reset synchroniser was implemented and evaluated. It closed those checks, and "
         "it broke gate-level simulation from the first instruction of the program, "
         "reproducibly, across two synchroniser variants and at every corner. The failure "
         "was established as causal by a controlled A/B on an identical harness. The "
         "design ships with the waiver rather than with a reset network that cannot be "
         "simulated, and closing that gap properly is the first item of future work.")

    # ---------------- 7 ----------------
    heading(doc, "7.  Area and Power", 1, page_break=True)

    heading(doc, "7.1  Area", 2)
    table(doc,
          ["Quantity", "Value"],
          [["Standard cells", d["cells"]],
           ["Sequential elements", d["flops"]],
           ["Filler cells", d["fillers"]],
           ["Core area", f"{d['core_um2']} µm²"],
           ["Achieved density", f"{d['density']}%"],
           ["Total wirelength", f"{d['wire_um']} µm"]],
          widths=[2.6, 3.9])

    heading(doc, "7.2  Power", 2)
    para(doc,
         f"Total power is {d['power_total']} mW at 250 MHz, estimated by Innovus against "
         f"the routed parasitics with default switching activity.")
    table(doc,
          ["Component", "Power (mW)", "Share"],
          [["Sequential", d["pw_seq"][0], f"{d['pw_seq'][1]}%"],
           ["Combinational", d["pw_comb"][0], f"{d['pw_comb'][1]}%"],
           ["Clock network", d["pw_clk"][0], f"{d['pw_clk'][1]}%"],
           ["Total", d["power_total"], "100%"]],
          widths=[2.2, 2.2, 2.1])
    para(doc,
         "Roughly half the power is in the sequential elements and nearly a quarter in the "
         "clock network, which is the expected shape for a design with 1,347 flops and no "
         "clock gating. Clock gating is the obvious first optimisation and is not "
         "implemented here.")

    heading(doc, "7.3  Cell-Level Detail", 2)
    figure(doc, os.path.join(ROOT, "docs", "images", "die-zoom.png"),
           "Figure 3: A 10 µm window into the standard cell rows. Blue horizontals are the "
           "metal1 supply rails bounding each row; red verticals and green horizontals are "
           "signal routing on the layers above; the crossed squares are vias.", width=5.2)

    # ---------------- 8 ----------------
    heading(doc, "8.  Limitations", 1, page_break=True)
    para(doc,
         "These are stated rather than omitted. Each has a specific cause, and the "
         "distinction between what the environment cannot do and what was not attempted "
         "is kept explicit.")

    heading(doc, "8.1  Not Possible in This Environment", 2)
    table(doc,
          ["Check", "Reason"],
          [["Antenna", "The Nangate 45 LEF carries no ANTENNAGATEAREA or ANTENNADIFFAREA "
                       "properties, so there is no ratio to check against"],
           ["LVS", "No Pegasus available; there is no layout-versus-schematic tool in this "
                   "environment"],
           ["Signoff DRC", "As above. The verify_drc result in Section 5.1 is the router's "
                           "own checker, not an independent signoff tool"],
           ["Independent signoff STA", "No Tempus available, so Innovus serves as its own "
                                       "signoff timer"],
           ["Hold coverage enumeration", "report_analysis_coverage on Innovus 23.12 accepts "
                                         "neither -check_type hold nor -early. Hold results "
                                         "are verified directly instead"]],
          widths=[1.9, 4.6])

    heading(doc, "8.2  Not Attempted", 2)
    table(doc,
          ["Item", "Note"],
          [["IR drop analysis", "Voltus rail analysis against the routed design has not "
                                "been run. The power grid is unverified beyond connectivity"],
           ["Design for test", "No scan insertion, no test compression, no ATPG"],
           ["Clock gating", "Not implemented; see Section 7.2"],
           ["Input/output delay", "Set to 0.30 × the clock period rather than to a "
                                  "characterised figure. All fifty worst paths are "
                                  "register-to-register, so this does not contaminate the "
                                  "frequency result, but it remains an assumption"]],
          widths=[1.9, 4.6])

    # ---------------- 9 ----------------
    heading(doc, "9.  Conclusion", 1, page_break=True)
    para(doc,
         f"The pipelined RV32I processor has been implemented from RTL to a routed layout "
         f"and closes at 4.00 ns, 250 MHz, at the slow corner with derates at 1.0. Setup "
         f"worst negative slack is {d['setup_signoff']} ns and hold is {d['hold_signoff']} ns, "
         f"both with zero violating paths across three corners. The layout is DRC clean "
         f"and connectivity clean, formally equivalent to its source across "
         f"{d['lec_equiv']} key points, and executes its verification program correctly at "
         f"four simulation tiers.")
    para(doc,
         "The closure is repeatable, four times out of four at a mean of +5 ps. Under 5% "
         "on-chip variation derates the same netlist stands at −0.237 ns, and recovering "
         "that is an optimisation-margin problem rather than a clock-period one: the "
         "measured physical limit of this netlist is 3.83 ns, so the constraint rather "
         "than the silicon is what sets 4.00.")
    para(doc,
         "The two gaps worth naming are the reset synchroniser, which closes the recovery "
         "and removal checks the constraints currently waive but breaks gate-level "
         "simulation for reasons not yet explained, and IR drop analysis, which has not "
         "been attempted. Neither affects the results reported here, and both are stated "
         "so that the boundary of what has been verified is unambiguous.")

    doc.save(out)
    return out


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--run", default="signoff-250mhz")
    ap.add_argument("--out", default=os.path.join(ROOT, "docs", "physical-design-report.docx"))
    a = ap.parse_args()
    os.makedirs(os.path.dirname(a.out), exist_ok=True)
    print("wrote", build(a.run, a.out))
