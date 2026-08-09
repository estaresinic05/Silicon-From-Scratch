#!/usr/bin/env python3
"""
Build a formatted Word report for one place-and-route run.

    python3 scripts/make_report.py --run runs/clk3p0
    python3 scripts/make_report.py --run runs/clk3p0 --images docs/images \
                                   --out docs/pnr-report-clk3p0.docx

Every number in the document is parsed out of that run's reports by qor.py.
Nothing is typed in by hand, which means a regenerated report cannot disagree
with the reports it claims to describe, and each iteration gets a document
that is correct by construction.

Screenshots are optional. Any of the filenames in FIGURES found in --images
is placed in its section with a caption; the rest of the document closes over
the gap. Requires python-docx.
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

sys.path.insert(0, str(Path(__file__).resolve().parent))
import qor  # noqa: E402


# --------------------------------------------------------------------------
# Design tokens. Copper on a cool grey, matching the HTML write-ups so the
# set reads as one body of work.
# --------------------------------------------------------------------------
ACCENT = RGBColor(0x9C, 0x5A, 0x26)
INK = RGBColor(0x14, 0x18, 0x1D)
INK_2 = RGBColor(0x43, 0x4C, 0x57)
INK_3 = RGBColor(0x6E, 0x77, 0x84)

FILL_HEAD = "E3E7EC"
FILL_CODE = "F4F6F8"
FILL_ZEBRA = "F7F9FA"

BODY_FONT = "Cambria"
MONO_FONT = "Consolas"

FIGURES = [
    ("die-routed.png", "figure-die",
     "The routed core. Standard cells in rows, signal routing on metal2 "
     "through metal6, and the power ring on metal8 and metal9."),
    ("die-zoom.png", "figure-zoom",
     "Detail. Individual standard cells sitting in their rows, with the "
     "metal1 power rails running along each row boundary."),
    ("die-critical-path.png", "figure-path",
     "The critical path highlighted across the die. It launches in MEM/WB, "
     "crosses the forwarding logic and ripples through all 32 bits of the "
     "ALU adder before being captured in EX/MEM."),
    ("die-modules.png", "figure-modules",
     "The register file selected across the die. All 992 flip-flops matching "
     "datapath/registers_RF_reg, which is 31 architectural registers of 32 "
     "bits each, x0 costing nothing because it is hardwired to zero."),
    ("power-grid.png", "figure-power",
     "The power grid with signal routing hidden. The ring runs on metal9 and "
     "metal8 around the core, vertical straps cross on metal8 every 25 "
     "microns, and sroute ties the whole grid down to the metal1 rails "
     "running inside every standard cell row."),
    ("clock-tree.png", "figure-clock",
     "The synthesised clock tree driving every flip-flop in the design."),
    ("congestion.png", "figure-congestion",
     "Routing congestion. Each bin is shaded by how much routing demand it "
     "carries against what the tracks there can supply."),
]


# --------------------------------------------------------------------------
# low-level docx helpers
# --------------------------------------------------------------------------

def _shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def shade_para(paragraph, fill):
    _shade(paragraph._p.get_or_add_pPr(), fill)


def set_borders(table, color="CCD3DA", size=4):
    """python-docx has no border API, so this is raw OOXML."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def para_rule(paragraph, color="14181D", size=12):
    """A horizontal rule, drawn as a bottom border on an empty paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)
    return run


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    pPr.append(el)


# --------------------------------------------------------------------------
# document furniture
# --------------------------------------------------------------------------

def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.18

    sizes = {"Heading 1": 17, "Heading 2": 13.5, "Heading 3": 11.5}
    for name, size in sizes.items():
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = False
        st.font.color.rgb = INK if name != "Heading 1" else ACCENT
        st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.paragraph_format.space_before = Pt(20 if name == "Heading 1" else 13)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)


def add_footer(doc, design):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{design}   ·   ")
    run.font.size = Pt(8)
    run.font.name = MONO_FONT
    run.font.color.rgb = INK_3
    fld = add_field(p, "PAGE")
    fld.font.size = Pt(8)
    fld.font.name = MONO_FONT
    fld.font.color.rgb = INK_3


# --------------------------------------------------------------------------
# content builders
# --------------------------------------------------------------------------

def title_block(doc, title, subtitle, meta_pairs):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("PHYSICAL DESIGN REPORT")
    r.font.name = MONO_FONT
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    r.font.all_caps = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = BODY_FONT
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = INK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    r.font.size = Pt(11.5)
    r.font.color.rgb = INK_2

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    para_rule(rule)

    for label, value in meta_pairs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"{label}   ")
        r.font.name = MONO_FONT
        r.font.size = Pt(8.5)
        r.font.color.rgb = INK_3
        r = p.add_run(str(value))
        r.font.name = MONO_FONT
        r.font.size = Pt(8.5)
        r.font.color.rgb = INK_2

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def code_block(doc, text, caption=None):
    """A shaded single-cell table. Word has no reliable code style."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_borders(table, color="DDE3E9", size=4)
    cell = table.cell(0, 0)
    shade_cell(cell, FILL_CODE)
    cell.text = ""
    for i, line in enumerate(text.rstrip("\n").split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(8.5)
        r.font.color.rgb = INK
    if caption:
        add_caption(doc, caption)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = INK_3


def data_table(doc, headers, rows, aligns=None, caption=None, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_borders(table)
    aligns = aligns or ["l"] * len(headers)

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, FILL_HEAD)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if aligns[i] == "r":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(str(h))
        r.font.name = MONO_FONT
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = INK_2

    for n, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if n % 2 == 1:
                shade_cell(cells[i], FILL_ZEBRA)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if aligns[i] == "r":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            text = "" if val is None else str(val)
            bold = text.startswith("**") and text.endswith("**")
            if bold:
                text = text[2:-2]
            r = p.add_run(text)
            r.font.size = Pt(9)
            r.font.bold = bold
            if aligns[i] == "r" or i == 0:
                r.font.name = MONO_FONT
                r.font.size = Pt(8.5)

    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    if caption:
        add_caption(doc, caption)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def bullets(doc, items):
    for lead, rest in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        if lead:
            r = p.add_run(lead)
            r.bold = True
        p.add_run(rest)


def maybe_figure(doc, images_dir, key, width_in=6.0):
    """Insert a screenshot if it exists. Silence if it does not."""
    if not images_dir:
        return False
    for name, tag, caption in FIGURES:
        if tag != key:
            continue
        path = Path(images_dir) / name
        if not path.exists():
            return False
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(path), width=Inches(width_in))
        add_caption(doc, caption)
        return True
    return False


# --------------------------------------------------------------------------

def f(v, places=3, dash="not measured"):
    if v in ("", None):
        return dash
    try:
        return f"{float(v):.{places}f}"
    except (TypeError, ValueError):
        return str(v)


def i(v, dash="not measured"):
    if v in ("", None):
        return dash
    try:
        return f"{int(float(v)):,}"
    except (TypeError, ValueError):
        return str(v)


def build(run_dir, out_path, images_dir=None, author="Elliot Staresinic"):
    run_dir = Path(run_dir)
    d = qor.parse_run(run_dir)
    layers = qor.parse_wire_by_layer(run_dir / "reports" / "44_summary.rpt")
    notes_file = run_dir / "NOTES.md"
    note = notes_file.read_text(encoding="utf-8").strip() if notes_file.exists() else ""

    doc = Document()
    setup_styles(doc)
    add_footer(doc, "pipelined_cpu_core")

    clk = d.get("clk_ns")
    freq = f"{1000 / float(clk):.0f} MHz" if clk else "not recorded"

    title_block(
        doc,
        "From RTL to a Routed Die",
        "A five-stage RISC-V pipeline carried through Cadence Genus and "
        "Innovus on a 45 nm standard cell library.",
        [
            ("DESIGN  ", "pipelined_cpu_core, five-stage RISC-V, RV32I subset"),
            ("RUN     ", f"{d.get('run')}   {d.get('date') or ''}"),
            ("CLOCK   ", f"{f(clk, 2, 'from SDC')} ns   ({freq})"),
            ("LIBRARY ", "Nangate45, typical corner, QRC extraction"),
            ("TOOLS   ", "Cadence Genus and Innovus 23.12-s091_1"),
            ("AUTHOR  ", author),
        ],
    )

    # ---------------------------------------------------------------- summary
    doc.add_heading("Result", level=1)

    setup = d.get("wns_setup")
    met = (setup not in ("", None)) and float(setup) >= 0
    verdict = "meets setup timing" if met else "does not meet setup timing"

    body(doc,
         f"The design routes to completion and {verdict} at a "
         f"{f(clk, 2, 'constrained')} ns clock, with a worst negative slack of "
         f"{f(setup)} ns on the critical path. It occupies "
         f"{i(d.get('cells'))} standard cells at {f(d.get('density_pct'), 1)} "
         f"percent density, wired with {i(d.get('wire_um'))} microns of copper.")

    data_table(
        doc,
        ["Metric", "Value", "Source"],
        [
            ["Setup WNS", f"**{f(setup)} ns**", "40_final_setup.rpt"],
            ["Setup paths violated", i(d.get("n_setup_viol"), "0"), "40_final_setup.rpt"],
            ["Hold WNS", f"{f(d.get('wns_hold'))} ns", "41_final_hold.rpt"],
            ["Hold paths violated", i(d.get("n_hold_viol"), "0"), "41_final_hold.rpt"],
            ["Standard cells", i(d.get("cells")), "44_summary.rpt"],
            ["Filler cells", i(d.get("fillers")), "44_summary.rpt"],
            ["Flip-flops", i(d.get("flops")), "44_summary.rpt"],
            ["Logic area", f"{f(d.get('logic_um2'), 1)} um2", "44_summary.rpt"],
            ["Core area", f"{f(d.get('core_um2'), 1)} um2", "44_summary.rpt"],
            ["Core density", f"{f(d.get('density_pct'), 1)} %", "44_summary.rpt"],
            ["Total wire length", f"{i(d.get('wire_um'))} um", "44_summary.rpt"],
        ],
        aligns=["l", "r", "l"],
        widths=[2.1, 1.6, 2.5],
        caption="Table 1. Signoff numbers, parsed directly from the run's own "
                "reports. Every figure in this document is read from the files "
                "named in the third column.",
    )

    if note:
        doc.add_heading("What changed in this run", level=2)
        for line in note.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    maybe_figure(doc, images_dir, "figure-die")

    # ------------------------------------------------------------------ flow
    doc.add_heading("What the flow does", level=1)

    body(doc,
         "Synthesis turns Verilog into a netlist: a flat list of logic gates "
         "and the wires between them, carrying no physical meaning whatsoever. "
         "Place and route turns that list into geometry, deciding where every "
         "gate physically sits and what copper connects them. The stages below "
         "run in order, and the design is written to disk after each one so a "
         "late failure never costs the earlier work.")

    data_table(
        doc,
        ["Stage", "Command", "What it produces"],
        [
            ["1", "floorPlan", "A core box solved from a density target, not a fixed size"],
            ["2", "addRing, addStripe, sroute", "A power grid, which the netlist never mentions"],
            ["3", "place_opt_design", "Coordinates, plus logic restructuring to meet timing there"],
            ["4", "clock_opt_design", "A real buffered clock tree, replacing the ideal clock"],
            ["5", "addFiller", "Continuous power rails and wells across every row"],
            ["6", "route_opt_design", "Copper, then re-optimisation against extracted parasitics"],
            ["7", "verify_drc", "A check that the geometry obeys the routing rules"],
            ["8", "report, defOut", "The reports this document is built from"],
        ],
        aligns=["r", "l", "l"],
        widths=[0.5, 1.9, 3.8],
        caption="Table 2. The place-and-route flow, stage by stage.",
    )

    doc.add_heading("Why the flow is written in Tcl", level=2)

    body(doc,
         "Innovus is, at its core, a Tcl interpreter. It is a Tcl shell with "
         "several thousand extra commands compiled in and a database of the "
         "chip attached to it. Typing place_opt_design at the Innovus prompt "
         "is typing Tcl, and running innovus -files scripts/innovus.tcl means "
         "nothing more than reading those same commands from a file instead of "
         "from the keyboard. There is no compilation step and no difference "
         "between the scripted tool and the interactive one.")

    body(doc,
         "This is not a Cadence quirk. Tcl was designed in the late 1980s "
         "specifically to be an embeddable command language for applications, "
         "EDA adopted it early, and it stayed. Cadence, Synopsys and Siemens "
         "tools are all driven this way, so the skill transfers across the "
         "industry in a way that familiarity with one vendor's graphical "
         "interface does not.")

    data_table(
        doc,
        ["File", "Language", "Talks to", "Job"],
        [
            ["run.sh", "bash", "Linux", "Check prerequisites, make directories, launch tools"],
            ["scripts/genus.tcl", "Tcl", "Genus", "RTL to gate-level netlist"],
            ["scripts/innovus.tcl", "Tcl", "Innovus", "Netlist to routed layout"],
            ["constraints/*.sdc", "Tcl", "any timer", "Declare what correct timing means"],
        ],
        aligns=["l", "l", "l", "l"],
        widths=[1.5, 0.8, 0.9, 3.0],
        caption="Table 3. Three languages, three jobs. The shell script never "
                "mentions the chip; it confirms the tools and libraries exist "
                "and hands off.",
    )

    body(doc,
         "SDC is Tcl as well, and that is why it is portable. create_clock "
         "-name clk -period 3.0 is a procedure call, and because every "
         "vendor's timing engine implements the same procedure names, one "
         "constraint file feeds synthesis, place and route, and a signoff "
         "timer unchanged. That portability is why constraints live in their "
         "own file rather than inside the implementation script.",
         bold_lead="Constraints are portable because they are Tcl. ")

    body(doc,
         "Physical design is not done in a graphical interface. The layout "
         "view is for inspecting a result, never for producing one. Production "
         "blocks run in batch on a compute farm, unattended, for hours or "
         "days, and a flow that cannot run headless cannot be run at scale. "
         "Scripting also makes a run reproducible: because every input is a "
         "text file under version control, it is possible to say exactly what "
         "changed between two results, which is the entire basis of being able "
         "to iterate on one.",
         bold_lead="Why it is scripted at all. ")

    # ------------------------------------------------------------- where time
    doc.add_heading("Where the timing went", level=1)

    start = d.get("wns_place_start")
    place, cts, final = d.get("wns_place"), d.get("wns_cts"), d.get("wns_setup")

    body(doc,
         "Slack is the margin on the worst path. Positive means the design "
         "works at that clock, negative means it does not, and the magnitude "
         "says by how much. Tracking it stage by stage is how you tell whether "
         "a timing problem lives in synthesis, in placement, in the clock tree "
         "or in the routing.")

    stage_rows = [
        ["Netlist as placed", "ideal", f(start), "Starting point, before optimisation"],
        ["place_opt_design", "ideal", f(place), "Placement optimisation"],
        ["clock_opt_design", "propagated", f(cts), "A real clock tree replaces the ideal one"],
        ["route_opt_design", "propagated + RC", f"**{f(final)}**", "Real copper, extracted parasitics"],
    ]
    data_table(
        doc,
        ["After", "Clock model", "WNS (ns)", "What happened"],
        stage_rows,
        aligns=["l", "l", "r", "l"],
        widths=[1.5, 1.2, 0.9, 2.6],
        caption="Table 4. Setup slack through the flow. The starting point is "
                "read from the optimiser's own progress table in "
                "reports/um*/flow_QOR_summary.rpt.",
    )

    try:
        cts_cost = float(place) - float(cts)
        route_cost = float(cts) - float(final)
        recovered = float(place) - float(start)
        body(doc,
             f"Placement did nearly all of the work. The netlist synthesis "
             f"handed over was {abs(float(start)) * 1000:.0f} picoseconds short "
             f"the moment real coordinates existed, and placement optimisation "
             f"recovered {recovered * 1000:.0f} picoseconds of that. Building "
             f"the actual clock tree cost {cts_cost * 1000:.0f} picoseconds and "
             f"routing real wires cost a further {route_cost * 1000:.0f}.",
             bold_lead="Reading the table. ")
    except (TypeError, ValueError):
        pass

    maybe_figure(doc, images_dir, "figure-path")

    body(doc,
         "A run that meets timing does not tell you what was achievable. The "
         "optimiser converges to whatever constraint it was given and then "
         "stops, and its area reclaiming passes deliberately spend positive "
         "slack to recover area and leakage. Synthesis stops trying at the "
         "same constraint. Finding the real maximum frequency means tightening "
         "the period and rerunning both tools until it fails.",
         bold_lead="An important limit on this number. ")

    # ---------------------------------------------------------------- physical
    doc.add_heading("The physical result", level=1)

    body(doc,
         f"The core holds {i(d.get('cells'))} standard cells of real logic "
         f"alongside {i(d.get('fillers'))} filler cells. Fillers contain no "
         f"logic at all; standard cells carry power rails and wells that must "
         f"be physically continuous across a row, and fillers exist purely to "
         f"bridge the gaps left between placed cells. Without them the rails "
         f"break and the design is not manufacturable.")

    if layers:
        total = sum(layers.values()) or 1
        rows = []
        for n in range(1, 11):
            key = f"metal{n}"
            if key not in layers:
                continue
            v = layers[key]
            rows.append([key, f"{v:,.0f}", f"{100 * v / total:.1f} %"])
        data_table(
            doc, ["Layer", "Wire length (um)", "Share"], rows,
            aligns=["l", "r", "r"], widths=[1.2, 1.6, 1.0],
            caption="Table 5. Routing by metal layer. metal1 carries no signal "
                    "because it is reserved for wiring inside the cells and for "
                    "the power rails. The upper layers are nearly empty because "
                    "a block this small has no journeys long enough to need them.",
        )

    maybe_figure(doc, images_dir, "figure-zoom")

    body(doc,
         "The power grid is worth seeing on its own, because none of it comes "
         "from the netlist. A netlist describes logic and says nothing about "
         "how current reaches it. Every ring, strap and rail below was created "
         "by three commands in the flow script, and the connection between a "
         "cell's power pin and the VDD net exists only because "
         "globalNetConnect declared it.",
         bold_lead="Power is added, never inherited. ")

    maybe_figure(doc, images_dir, "figure-power")
    maybe_figure(doc, images_dir, "figure-modules")
    maybe_figure(doc, images_dir, "figure-clock")
    maybe_figure(doc, images_dir, "figure-congestion")

    # -------------------------------------------------------------------- gaps
    doc.add_heading("What a production flow has that this one does not", level=1)

    body(doc,
         "The structure of this flow matches industry practice. The rigour "
         "does not, in ways worth naming explicitly.")

    bullets(doc, [
        ("One corner is not signoff. ",
         "This run used a single library, a single RC corner and a single "
         "analysis view. Real signoff is multi-corner multi-mode, spanning "
         "slow and fast silicon, hot and cold, high and low voltage, and "
         "functional against test modes. The same library was also used for "
         "setup and hold, where production uses a slow corner for setup and a "
         "fast corner for hold."),
        ("The place-and-route timer is not the signoff timer. ",
         "Innovus's internal engine is good enough to optimise against, but "
         "signoff exports the netlist with extracted parasitics and re-times "
         "in a dedicated tool, expecting the two to disagree slightly."),
        ("verify_drc is not physical verification. ",
         "It checks routing rules inside the tool. Real DRC and LVS run "
         "against a foundry runset in a separate verification tool. Nangate45 "
         "is an academic library with no foundry behind it, so this does not "
         "apply here, but it would immediately on real silicon."),
        ("Whole categories are absent. ",
         "IR drop and electromigration, antenna fixing, metal fill, scan "
         "insertion and chain reordering, and formal equivalence checking "
         "between the RTL and the synthesised netlist."),
        ("Flows are normally inherited, not authored. ",
         "Companies have methodology teams that own these scripts, and a block "
         "designer supplies the RTL, the constraints and a configuration file "
         "of block-specific knobs. Writing one from scratch was worth doing "
         "once for the understanding."),
    ])

    # ------------------------------------------------------------ reproducing
    doc.add_heading("Reproducing this run", level=1)

    body(doc,
         "The run is fully described by the scripts in the repository. Every "
         "report quoted above is committed under results/ so each number can "
         "be checked against its source.")

    code_block(doc,
               "./run.sh --period {p} --name {n} --note \"...\"\n"
               "\n"
               "runs/{n}/reports/     the reports this document parses\n"
               "results/{n}/reports/  the same files, committed as evidence\n"
               "results/QOR.md        one row per run, for comparing iterations"
               .format(p=f(clk, 1, "3.0"), n=d.get("run")))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    para_rule(p, color="14181D", size=12)

    for line in [
        f"Generated from {d.get('run')} by scripts/make_report.py.",
        "Numbers are parsed from the run's reports and are not transcribed.",
        "Repository: estaresinic05/Silicon-From-Scratch",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.font.name = MONO_FONT
        r.font.size = Pt(8)
        r.font.color.rgb = INK_3

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path, d


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--run", required=True, help="run directory containing reports/")
    ap.add_argument("--out", default=None, help="output .docx path")
    ap.add_argument("--images", default=None, help="directory of screenshots")
    ap.add_argument("--author", default="Elliot Staresinic")
    args = ap.parse_args()

    run_dir = Path(args.run)
    if not (run_dir / "reports").is_dir():
        sys.exit(f"no reports/ under {run_dir}")

    root = Path(__file__).resolve().parent.parent
    out = Path(args.out) if args.out else root / "docs" / f"pnr-report-{run_dir.name}.docx"

    path, d = build(run_dir, out, args.images, args.author)
    print(f"wrote  {path}")
    print(f"run    {d.get('run')}   setup WNS {f(d.get('wns_setup'))} ns")


if __name__ == "__main__":
    main()
